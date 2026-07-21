"""Human-owned document library, sequential reader, and derived RAG index.

The imported bytes and extracted text are canonical.  Chunk vectors are a
regenerable sidecar, exactly like memory vectors: useful when healthy, never a
second source of truth.  Reading position is persona-specific so two members
of the household can inhabit the same human-owned source differently without
copying or rewriting it.
"""
from __future__ import annotations

import base64
import binascii
import hashlib
import io
import json
import math
import os
import re
import shutil
import tempfile
import time
from html.parser import HTMLParser
from pathlib import Path

from core.users import slugify


MAX_DOCUMENT_BYTES = 25 * 1024 * 1024
MAX_EXTRACTED_CHARS = 12_000_000
DOCUMENT_CONTEXT_BUDGET = 900
TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".rst", ".csv", ".tsv", ".json",
    ".yaml", ".yml", ".html", ".htm",
}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {".pdf", ".docx"}


class DocumentError(ValueError):
    pass


def private_document_access(speaker: str, local_human: str,
                            channel: str) -> tuple[bool, str]:
    """Private-by-default disclosure gate for prompt assembly."""
    if str(channel or "") == "room":
        return False, "room_channel_private_default"
    if str(speaker or "") != str(local_human or ""):
        return False, "speaker_is_not_local_owner"
    return True, "local_owner_private_turn"


class _HTMLText(HTMLParser):
    BLOCKS = {
        "address", "article", "aside", "blockquote", "br", "div", "dl",
        "dt", "dd", "figcaption", "figure", "footer", "h1", "h2", "h3",
        "h4", "h5", "h6", "header", "hr", "li", "main", "nav", "ol",
        "p", "pre", "section", "table", "tr", "ul",
    }

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.parts = []
        self.hidden = 0

    def handle_starttag(self, tag, attrs):
        if tag in {"script", "style", "noscript"}:
            self.hidden += 1
        elif tag in self.BLOCKS:
            self.parts.append("\n")

    def handle_endtag(self, tag):
        if tag in {"script", "style", "noscript"} and self.hidden:
            self.hidden -= 1
        elif tag in self.BLOCKS:
            self.parts.append("\n")

    def handle_data(self, data):
        if not self.hidden:
            self.parts.append(data)

    def text(self):
        return re.sub(r"\n\s*\n\s*\n+", "\n\n", "".join(self.parts)).strip()


def _atomic_json(path: Path, value) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    rendered = json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True)
    json.loads(rendered)
    fd, tmp = tempfile.mkstemp(prefix=".jnsq-doc-", suffix=".tmp",
                               dir=path.parent, text=True)
    try:
        with os.fdopen(fd, "w", encoding="utf-8", newline="\n") as handle:
            handle.write(rendered)
            handle.write("\n")
        os.replace(tmp, path)
    finally:
        if os.path.exists(tmp):
            os.unlink(tmp)


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "cp1252"):
        try:
            text = data.decode(encoding)
            if "\x00" not in text:
                return text
        except UnicodeDecodeError:
            continue
    raise DocumentError("document text encoding is not supported")


def _extract(data: bytes, suffix: str) -> tuple[str, str]:
    if suffix in TEXT_EXTENSIONS:
        text = _decode_text(data)
        if suffix in {".html", ".htm"}:
            parser = _HTMLText()
            parser.feed(text)
            text = parser.text()
            extractor = "html.parser"
        else:
            extractor = "decoded_text"
    elif suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise DocumentError(
                "PDF import needs pypdf; install the JNSQ requirements") from exc
        try:
            pages = []
            for number, page in enumerate(PdfReader(io.BytesIO(data)).pages, 1):
                pages.append(f"[Page {number}]\n{page.extract_text() or ''}")
            text = "\n\n".join(pages)
        except Exception as exc:
            raise DocumentError(f"PDF text extraction failed: {exc}") from exc
        extractor = "pypdf"
    elif suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise DocumentError(
                "DOCX import needs python-docx; install the JNSQ requirements") from exc
        try:
            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                parts.extend("\t".join(cell.text for cell in row.cells)
                             for row in table.rows)
            text = "\n\n".join(parts)
        except Exception as exc:
            raise DocumentError(f"DOCX text extraction failed: {exc}") from exc
        extractor = "python-docx"
    else:
        raise DocumentError(
            f"unsupported document type {suffix or '(none)'}; supported: "
            + ", ".join(sorted(SUPPORTED_EXTENSIONS)))
    text = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not text:
        raise DocumentError("document contains no extractable text")
    if len(text) > MAX_EXTRACTED_CHARS:
        raise DocumentError(
            f"extracted text exceeds the {MAX_EXTRACTED_CHARS:,}-character boundary")
    return text, extractor


def _structural_units(text: str) -> list[tuple[int, int]]:
    units = []
    for match in re.finditer(r"\S(?:.*?\S)?(?=\n\s*\n|\Z)", text, re.S):
        units.append((match.start(), match.end()))
    return units or [(0, len(text))]


def _sentence_units(text: str, start: int, end: int) -> list[tuple[int, int]]:
    found = []
    for match in re.finditer(r".*?(?:[.!?](?=\s)|\n|\Z)", text[start:end], re.S):
        left, right = start + match.start(), start + match.end()
        while left < right and text[left].isspace():
            left += 1
        while right > left and text[right - 1].isspace():
            right -= 1
        if right > left:
            found.append((left, right))
    return found or [(start, end)]


def _chunk_text(text: str) -> tuple[list[dict], dict]:
    # The target grows with the square root of the document and remains within
    # the inherited prompt-context envelope.  Long documents get broader
    # sections without turning a reader position into an arbitrary fixed grid.
    target = max(900, min(2400, round(math.sqrt(len(text)) * 16)))
    lower = round(target * 0.55)
    upper = round(target * 1.35)
    units = []
    for start, end in _structural_units(text):
        if end - start <= upper:
            units.append((start, end))
        else:
            units.extend(_sentence_units(text, start, end))

    chunks, current = [], []

    def emit(spans):
        if not spans:
            return
        start, end = spans[0][0], spans[-1][1]
        while start < end:
            hard_end = min(end, start + upper)
            if hard_end < end:
                boundary = text.rfind(" ", start + lower, hard_end)
                hard_end = boundary if boundary > start else hard_end
            value = text[start:hard_end].strip()
            if value:
                real_start = start + len(text[start:hard_end]) - len(
                    text[start:hard_end].lstrip())
                real_end = hard_end - len(text[start:hard_end]) + len(
                    text[start:hard_end].rstrip())
                chunks.append({"index": len(chunks), "text": value,
                               "char_start": real_start,
                               "char_end": real_end})
            start = hard_end
            while start < end and text[start].isspace():
                start += 1

    for span in units:
        if not current:
            current = [span]
            continue
        proposed = span[1] - current[0][0]
        current_size = current[-1][1] - current[0][0]
        if proposed > target and current_size >= lower:
            emit(current)
            current = [span]
        else:
            current.append(span)
    emit(current)
    return chunks, {"strategy": "structural_sqrt_v1", "target_chars": target,
                    "lower_chars": lower, "upper_chars": upper}


def _title_for(filename: str, text: str) -> str:
    heading = re.search(r"(?m)^#{1,6}\s+(.+?)\s*$", text[:12000])
    if heading:
        return heading.group(1).strip()[:200]
    return Path(filename).stem[:200] or "Untitled document"


def _clip(text: str, maximum: int) -> str:
    if len(text) <= maximum:
        return text
    cut = text.rfind(" ", 0, maximum)
    return text[:cut if cut > maximum // 2 else maximum].rstrip() + "\n[…excerpt continues]"


class DocumentLibrary:
    def __init__(self, repo: str | os.PathLike[str], user_id: str,
                 persona: str, *, now_fn=time.time):
        self.repo = Path(repo).resolve()
        self.user_id = slugify(user_id)
        persona = str(persona or "").strip()
        if not persona or Path(persona).name != persona or persona in {".", ".."}:
            raise DocumentError("persona name is outside the document-reader boundary")
        self.persona = persona
        self.root = self.repo / "users" / self.user_id / "documents"
        self.reader_root = (self.repo / "personas" / persona / "body" /
                            "document_reader")
        self.state_path = self.reader_root / "state.json"
        self.events_path = self.reader_root / "events.jsonl"
        self.receipts_path = self.reader_root / "receipts.jsonl"
        self.arc_path = self.reader_root / "reading_arc.json"
        self.notebook_path = self.reader_root / "notebook.jsonl"
        self.reports_root = self.reader_root / "reports"
        self.now_fn = now_fn

    def _append(self, path: Path, record: dict) -> dict:
        """Append one validated private reader record without rewriting history."""
        value = dict(record or {})
        value.setdefault("at", float(self.now_fn()))
        rendered = json.dumps(value, ensure_ascii=False, sort_keys=True)
        json.loads(rendered)
        path.parent.mkdir(parents=True, exist_ok=True)
        with path.open("a", encoding="utf-8", newline="\n") as handle:
            handle.write(rendered + "\n")
        return value

    def reader_events(self, kind: str = "", limit: int = 1000) -> list[dict]:
        limit = max(0, min(int(limit), 5000))
        try:
            lines = self.events_path.read_text(encoding="utf-8").splitlines()
        except OSError:
            return []
        records = []
        for line in lines[-limit:] if limit else ():
            try:
                value = json.loads(line)
            except (TypeError, ValueError):
                continue
            if isinstance(value, dict) and (not kind or value.get("kind") == kind):
                records.append(value)
        return records

    def record_receipt(self, record: dict) -> dict:
        return self._append(self.receipts_path, {
            "schema": 1, "persona": self.persona,
            "user_id": self.user_id, **dict(record or {})})

    def record_turn_exposure(self, anchors, *, exposure_id: str,
                             active_anchor: str = "",
                             retrieved_anchors=(), evidence: str =
                             "prompt_rendered") -> dict | None:
        """Record source sections that reached a completed turn-model call."""
        exposure_id = str(exposure_id or "").strip()[:160]
        if not exposure_id:
            raise DocumentError("document exposure id is required")
        existing = next((event for event in self.reader_events(
            "document_turn_exposure", 5000)
            if event.get("exposure_id") == exposure_id), None)
        if existing:
            return existing
        valid = []
        for anchor in dict.fromkeys(str(value or "").strip()
                                    for value in anchors or ()):
            if not anchor:
                continue
            valid.append(self.inspect_anchor(anchor, maximum=1)["anchor"])
        if not valid:
            return None
        active_anchor = str(active_anchor or "").strip()
        if active_anchor and active_anchor not in valid:
            active_anchor = ""
        retrieved = [str(value or "").strip()
                     for value in retrieved_anchors or ()]
        record = self._append(self.events_path, {
            "schema": 1, "kind": "document_turn_exposure",
            "persona": self.persona, "user_id": self.user_id,
            "exposure_id": exposure_id, "anchors": valid,
            "active_anchor": active_anchor,
            "retrieved_anchors": [anchor for anchor in retrieved
                                  if anchor in valid],
            "evidence": str(evidence or "prompt_rendered")[:80],
        })
        self._reconcile_reading_arc()
        return record

    def import_turn_exposure_history(self, memories) -> int:
        """Index old turn-anchor receipts without changing source or memory."""
        existing = {str(event.get("exposure_id") or "")
                    for event in self.reader_events(
                        "document_turn_exposure", 5000)}
        imported = 0
        for memory in memories or ():
            if not isinstance(memory, dict) or memory.get("type") != "turn":
                continue
            fields = memory.get("fields") or {}
            anchors = list(fields.get("document_anchors") or ())
            if not anchors:
                continue
            exposure_id = str(fields.get("document_exposure_id")
                              or f"memory:{memory.get('id') or ''}")[:160]
            if not exposure_id or exposure_id in existing:
                continue
            valid = []
            for anchor in dict.fromkeys(str(value or "").strip()
                                        for value in anchors):
                try:
                    valid.append(self.inspect_anchor(anchor, maximum=1)["anchor"])
                except DocumentError:
                    continue
            if not valid:
                continue
            self._append(self.events_path, {
                "schema": 1, "kind": "document_turn_exposure",
                "persona": self.persona, "user_id": self.user_id,
                "exposure_id": exposure_id, "anchors": valid,
                "active_anchor": "", "retrieved_anchors": [],
                "evidence": ("prompt_rendered" if fields.get(
                    "document_exposure_id") else "historical_turn_anchor"),
                "source_timestamp": memory.get("timestamp"),
            })
            existing.add(exposure_id)
            imported += 1
        if imported:
            self._reconcile_reading_arc()
        return imported

    def _doc_dir(self, doc_id: str) -> Path:
        doc_id = str(doc_id or "")
        if not re.fullmatch(r"doc_[0-9a-f]{16}", doc_id):
            raise DocumentError("invalid document id")
        return self.root / doc_id

    def _metadata(self, doc_id: str) -> dict:
        path = self._doc_dir(doc_id) / "document.json"
        try:
            value = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, ValueError, TypeError) as exc:
            raise DocumentError("document does not exist or its metadata is invalid") from exc
        return value

    def _chunks(self, doc_id: str) -> list[dict]:
        path = self._doc_dir(doc_id) / "chunks.json"
        try:
            value = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, ValueError, TypeError) as exc:
            raise DocumentError("document chunks are unavailable") from exc
        if not isinstance(value, list):
            raise DocumentError("document chunk index is invalid")
        return value

    def list_documents(self) -> list[dict]:
        if not self.root.is_dir():
            return []
        found = []
        for path in sorted(self.root.glob("doc_*/document.json")):
            try:
                rec = json.loads(path.read_text(encoding="utf-8"))
            except (OSError, ValueError, TypeError):
                continue
            if isinstance(rec, dict) and rec.get("id") == path.parent.name:
                found.append(rec)
        return sorted(found, key=lambda rec: (
            -float(rec.get("imported_at", 0.0)), str(rec.get("title", ""))))

    def has_documents(self) -> bool:
        return bool(self.list_documents())

    def import_bytes(self, filename: str, data: bytes,
                     content_type: str = "") -> dict:
        filename = Path(str(filename or "")).name.strip()
        if not filename:
            raise DocumentError("document filename is required")
        if not isinstance(data, (bytes, bytearray)):
            raise DocumentError("document payload must be bytes")
        data = bytes(data)
        if not data:
            raise DocumentError("document payload is empty")
        if len(data) > MAX_DOCUMENT_BYTES:
            raise DocumentError(
                f"document exceeds the {MAX_DOCUMENT_BYTES // (1024 * 1024)} MB boundary")
        suffix = Path(filename).suffix.casefold()
        if suffix not in SUPPORTED_EXTENSIONS:
            raise DocumentError(
                "unsupported document type; supported: "
                + ", ".join(sorted(SUPPORTED_EXTENSIONS)))
        digest = hashlib.sha256(data).hexdigest()
        doc_id = f"doc_{digest[:16]}"
        target = self._doc_dir(doc_id)
        if target.is_dir():
            record = self._metadata(doc_id)
            return {**record, "duplicate": True}

        text, extractor = _extract(data, suffix)
        chunks, chunking = _chunk_text(text)
        vectors = None
        vector_receipt = {"status": "unavailable", "rows": 0}
        try:
            from core.memory_emotion.vectors import embed_texts
            vectors = embed_texts([chunk["text"] for chunk in chunks])
            if vectors is not None:
                vector_receipt = {"status": "healthy", "rows": len(vectors)}
        except Exception as exc:
            vector_receipt = {"status": "unavailable", "rows": 0,
                              "error_type": type(exc).__name__}

        imported_at = float(self.now_fn())
        record = {
            "schema": 1,
            "id": doc_id,
            "owner": self.user_id,
            "title": _title_for(filename, text),
            "filename": filename,
            "extension": suffix,
            "content_type": str(content_type or ""),
            "sha256": digest,
            "bytes": len(data),
            "characters": len(text),
            "chunk_count": len(chunks),
            "chunking": chunking,
            "extractor": extractor,
            "imported_at": imported_at,
            "vectors": vector_receipt,
        }
        self.root.mkdir(parents=True, exist_ok=True)
        staging = Path(tempfile.mkdtemp(prefix=".jnsq-doc-", dir=self.root))
        try:
            (staging / f"source{suffix}").write_bytes(data)
            (staging / "text.txt").write_text(text, encoding="utf-8", newline="\n")
            _atomic_json(staging / "chunks.json", chunks)
            _atomic_json(staging / "document.json", record)
            if vectors is not None:
                import numpy as np
                np.save(staging / "vectors.npy", vectors)
            try:
                os.replace(staging, target)
            except FileExistsError:
                return {**self._metadata(doc_id), "duplicate": True}
        finally:
            if staging.exists():
                shutil.rmtree(staging, ignore_errors=True)
        return {**record, "duplicate": False}

    def import_data_url(self, filename: str, data_url: str,
                        content_type: str = "") -> dict:
        prefix, separator, encoded = str(data_url or "").partition(",")
        if not separator or not prefix.startswith("data:") or ";base64" not in prefix:
            raise DocumentError("document payload must be a base64 data URL")
        if len(encoded) > (MAX_DOCUMENT_BYTES * 4 // 3) + 16:
            raise DocumentError(
                f"document exceeds the {MAX_DOCUMENT_BYTES // (1024 * 1024)} MB boundary")
        try:
            data = base64.b64decode(encoded, validate=True)
        except (binascii.Error, ValueError, TypeError) as exc:
            raise DocumentError("document payload is not valid base64") from exc
        declared = prefix[5:].split(";", 1)[0]
        return self.import_bytes(filename, data, content_type or declared)

    def _state(self) -> dict:
        try:
            value = json.loads(self.state_path.read_text(encoding="utf-8"))
        except (OSError, ValueError, TypeError):
            return {}
        if (not isinstance(value, dict)
                or value.get("user_id") != self.user_id):
            return {}
        return value

    def open(self, doc_id: str, position: int = 0) -> dict:
        chunks = self._chunks(doc_id)
        if not chunks:
            raise DocumentError("document has no readable chunks")
        position = int(position)
        if position < 0 or position >= len(chunks):
            raise DocumentError(
                f"position must be between 0 and {len(chunks) - 1}")
        _atomic_json(self.state_path, {
            "schema": 1, "user_id": self.user_id, "persona": self.persona,
            "doc_id": doc_id, "position": position,
            "updated_at": float(self.now_fn()),
        })
        return self.reader_status(include_text=True)

    def navigate(self, action: str, position: int | None = None) -> dict:
        state = self._state()
        if not state.get("doc_id"):
            raise DocumentError("no document is open")
        chunks = self._chunks(state["doc_id"])
        current = max(0, min(int(state.get("position", 0)), len(chunks) - 1))
        action = str(action or "").casefold()
        if action == "next":
            wanted = min(len(chunks) - 1, current + 1)
        elif action == "previous":
            wanted = max(0, current - 1)
        elif action == "jump":
            if position is None:
                raise DocumentError("jump navigation requires a position")
            wanted = int(position)
        else:
            raise DocumentError("navigation action must be next, previous, or jump")
        return self.open(state["doc_id"], wanted)

    def reader_status(self, include_text: bool = True) -> dict:
        state = self._state()
        doc_id = state.get("doc_id")
        if not doc_id:
            return {"active": False, "persona": self.persona,
                    "owner": self.user_id}
        try:
            metadata = self._metadata(doc_id)
            chunks = self._chunks(doc_id)
        except DocumentError:
            return {"active": False, "persona": self.persona,
                    "owner": self.user_id, "stale_reference": doc_id}
        position = max(0, min(int(state.get("position", 0)), len(chunks) - 1))
        chunk = dict(chunks[position])
        if not include_text:
            chunk.pop("text", None)
        anchor = f"{doc_id}#{position + 1}"
        return {
            "active": True, "persona": self.persona, "owner": self.user_id,
            "document": metadata, "position": position,
            "section": position + 1, "total": len(chunks),
            "progress": (position + 1) / len(chunks),
            "has_previous": position > 0,
            "has_next": position < len(chunks) - 1,
            "anchor": anchor, "chunk": chunk,
        }

    def inspect_anchor(self, anchor: str, *, maximum: int = 3200) -> dict:
        """Inspect one canonical section through an exact anchor only.

        This is the narrow door used by the writing desk.  It performs no
        search and accepts no path: the human must already have admitted the
        anchor, and the caller remains responsible for that admission check.
        """
        match = re.fullmatch(r"(doc_[0-9a-f]{16})#([1-9][0-9]*)",
                             str(anchor or "").strip())
        if not match:
            raise DocumentError("document anchor is invalid")
        doc_id, section = match.group(1), int(match.group(2))
        metadata = self._metadata(doc_id)
        chunks = self._chunks(doc_id)
        index = section - 1
        if index < 0 or index >= len(chunks):
            raise DocumentError("document anchor section does not exist")
        maximum = max(1, min(int(maximum), 12000))
        chunk = dict(chunks[index])
        text = str(chunk.get("text") or "")
        return {
            "anchor": f"{doc_id}#{section}",
            "doc_id": doc_id,
            "section": section,
            "total": len(chunks),
            "title": metadata.get("title"),
            "filename": metadata.get("filename"),
            "char_start": chunk.get("char_start"),
            "char_end": chunk.get("char_end"),
            "content": text[:maximum],
            "source_chars": len(text),
            "truncated": len(text) > maximum,
            "sha256": hashlib.sha256(text.encode("utf-8")).hexdigest(),
            "ownership": "human_owned_document",
        }

    def encounter(self, anchor: str, *, action: str, query: str = "",
                  report: str = "", why: str = "", run_id: str = "") -> dict:
        """Persist one chosen encounter; source text is never copied here."""
        inspected = self.inspect_anchor(anchor, maximum=1)
        action = str(action or "quiet").strip().casefold()
        if action not in {"quiet", "continue", "search", "bookmark", "report",
                          "pause"}:
            raise DocumentError("document encounter action is invalid")
        query = re.sub(r"\s+", " ", str(query or "")).strip()[:300]
        report = str(report or "").strip()[:8000]
        if action == "search" and len(query) < 2:
            action, query = "quiet", ""
        if action == "report" and not report:
            action = "quiet"
        if action != "search":
            query = ""
        if action != "report":
            report = ""
        record = self._append(self.events_path, {
            "schema": 1, "kind": "document_encounter",
            "persona": self.persona, "user_id": self.user_id,
            "anchor": inspected["anchor"], "doc_id": inspected["doc_id"],
            "section": inspected["section"], "action": action,
            "query": query, "why": str(why or "")[:500],
            "run_id": str(run_id or "")[:160],
        })
        self.open(inspected["doc_id"], inspected["section"] - 1)
        if action == "report":
            created = self.create_report(
                inspected["anchor"], report, run_id=run_id)
            record["report_id"] = created["report_id"]
            record["report_anchor"] = created["anchor"]
        return record

    def create_report(self, source_anchor: str, content: str, *,
                      run_id: str = "") -> dict:
        source = self.inspect_anchor(source_anchor, maximum=1)
        content = str(content or "").strip()
        if not content:
            raise DocumentError("document report content is required")
        if f"[{source['anchor']}]" not in content:
            content += f"\n\nSource: [{source['anchor']}]"
        digest = hashlib.sha256(content.encode("utf-8")).hexdigest()
        report_id = f"drep_{digest[:16]}"
        path = self.reports_root / f"{report_id}.json"
        record = {
            "schema": 1, "report_id": report_id,
            "anchor": f"{report_id}#1", "persona": self.persona,
            "user_id": self.user_id, "source_anchor": source["anchor"],
            "title": f"Reading report: {source.get('title') or source['anchor']}",
            "content": content, "sha256": digest,
            "created_at": float(self.now_fn()), "run_id": str(run_id or "")[:160],
        }
        if not path.exists():
            _atomic_json(path, record)
        self._append(self.events_path, {
            "schema": 1, "kind": "document_report_created",
            **{key: record[key] for key in (
                "report_id", "anchor", "source_anchor", "created_at", "run_id")},
        })
        return record

    def report(self, report_id: str) -> dict:
        report_id = str(report_id or "")
        if not re.fullmatch(r"drep_[0-9a-f]{16}", report_id):
            raise DocumentError("document report id is invalid")
        try:
            value = json.loads((self.reports_root / f"{report_id}.json").read_text(
                encoding="utf-8"))
        except (OSError, TypeError, ValueError) as exc:
            raise DocumentError("document report does not exist") from exc
        if value.get("report_id") != report_id:
            raise DocumentError("document report identity changed")
        return value

    def inspect_report_anchor(self, anchor: str, *, maximum: int = 5200) -> dict:
        match = re.fullmatch(r"(drep_[0-9a-f]{16})#1", str(anchor or ""))
        if not match:
            raise DocumentError("document report anchor is invalid")
        report = self.report(match.group(1))
        content = str(report.get("content") or "")
        if hashlib.sha256(content.encode("utf-8")).hexdigest() != report.get("sha256"):
            raise DocumentError("document report digest changed")
        maximum = max(1, min(int(maximum), 12000))
        return {
            "anchor": report["anchor"], "title": report["title"],
            "content": content[:maximum], "source_chars": len(content),
            "truncated": len(content) > maximum, "sha256": report["sha256"],
            "source_anchor": report["source_anchor"],
            "ownership": "persona_private_document_report",
        }

    def pending_reports(self) -> list[dict]:
        settled = {r.get("report_id") for r in self.reader_events(limit=5000)
                   if r.get("kind") in {"document_report_handed_off",
                                        "document_report_settled"}}
        reports = []
        for path in sorted(self.reports_root.glob("drep_*.json")) \
                if self.reports_root.is_dir() else ():
            try:
                value = json.loads(path.read_text(encoding="utf-8"))
            except (OSError, TypeError, ValueError):
                continue
            if value.get("report_id") not in settled:
                reports.append(value)
        return reports

    def settle_report(self, report_id: str, *, run_id: str = "") -> dict:
        report = self.report(report_id)
        existing = next((r for r in self.reader_events(limit=5000)
                         if r.get("report_id") == report_id and r.get("kind") in {
                             "document_report_settled",
                             "document_report_handed_off"}), None)
        if existing:
            return existing
        return self._append(self.events_path, {
            "schema": 1, "kind": "document_report_settled",
            "report_id": report_id, "anchor": report["anchor"],
            "source_anchor": report["source_anchor"],
            "run_id": str(run_id or "")[:160],
        })

    def mark_report_handed_off(self, report_id: str, *, seed_id: str,
                               run_id: str = "") -> dict:
        report = self.report(report_id)
        existing = next((r for r in self.reader_events(
            "document_report_handed_off", 5000)
            if r.get("report_id") == report_id), None)
        if existing:
            return existing
        return self._append(self.events_path, {
            "schema": 1, "kind": "document_report_handed_off",
            "report_id": report_id, "anchor": report["anchor"],
            "source_anchor": report["source_anchor"],
            "seed_id": str(seed_id or "")[:160],
            "run_id": str(run_id or "")[:160],
        })

    def notebook_entries(self, doc_id: str = "", limit: int = 5000) -> list[dict]:
        limit = max(0, min(int(limit), 5000))
        try:
            lines = self.notebook_path.read_text(encoding="utf-8").splitlines()
        except OSError:
            return []
        found = []
        for line in lines[-limit:] if limit else ():
            try:
                value = json.loads(line)
            except (TypeError, ValueError):
                continue
            if not isinstance(value, dict):
                continue
            if doc_id and value.get("doc_id") != doc_id:
                continue
            found.append(value)
        return found

    def record_notebook_entry(self, anchor: str, *, observation: str,
                              feelings=None, action: str = "quiet",
                              run_id: str = "") -> dict | None:
        inspected = self.inspect_anchor(anchor, maximum=1)
        observation = re.sub(r"\s+", " ", str(observation or "")).strip()[:800]
        feelings = {str(key)[:40]: max(0.0, min(1.0, float(value)))
                    for key, value in list(dict(feelings or {}).items())[:4]}
        if not observation and not feelings:
            return None
        run_id = str(run_id or "")[:160]
        if run_id:
            prior = next((entry for entry in self.notebook_entries(
                inspected["doc_id"], 5000) if entry.get("run_id") == run_id), None)
            if prior:
                return prior
        return self._append(self.notebook_path, {
            "schema": 1, "kind": "document_reading_note",
            "persona": self.persona, "user_id": self.user_id,
            "doc_id": inspected["doc_id"], "anchor": inspected["anchor"],
            "section": inspected["section"], "observation": observation,
            "feelings": feelings, "action": str(action or "quiet")[:40],
            "run_id": run_id,
        })

    def notebook_context(self, doc_id: str, maximum: int = 2200) -> str:
        """Render the recent tail by pressure budget, never by arbitrary count."""
        maximum = max(0, min(int(maximum), 6000))
        if not maximum:
            return ""
        chosen, used = [], 0
        for entry in reversed(self.notebook_entries(doc_id, 5000)):
            observation = str(entry.get("observation") or "").strip()
            feelings = ", ".join(sorted((entry.get("feelings") or {}).keys()))
            detail = observation or (f"felt {feelings}" if feelings else "encountered")
            line = f"[{entry.get('anchor')}] {detail}"
            cost = len(line) + 1
            if chosen and used + cost > maximum:
                break
            chosen.append(line[:maximum] if not chosen else line)
            used += cost
        return "\n".join(reversed(chosen))[:maximum]

    def reading_coverage(self, doc_id: str) -> dict:
        metadata = self._metadata(doc_id)
        total = int(metadata.get("chunk_count") or 0)
        autonomous, conversation, inferred = set(), set(), set()
        latest = None
        for event in self.reader_events(limit=5000):
            kind = event.get("kind")
            if kind == "document_encounter" and event.get("doc_id") == doc_id:
                anchor = str(event.get("anchor") or "")
                autonomous.add(anchor)
                latest = {"route": "autonomous", **event}
            elif kind == "document_turn_exposure":
                for anchor in event.get("anchors") or ():
                    if str(anchor).startswith(doc_id + "#"):
                        conversation.add(str(anchor))
                        if event.get("evidence") == "historical_turn_anchor":
                            inferred.add(str(anchor))
                        latest = {"route": "conversation", **event,
                                  "anchor": str(anchor)}
        known = autonomous | conversation
        def sections(values):
            return sorted(int(anchor.rsplit("#", 1)[1]) for anchor in values)
        return {
            "doc_id": doc_id, "title": metadata.get("title"), "total": total,
            "known_count": len(known),
            "coverage": (len(known) / total if total else 0.0),
            "autonomous_count": len(autonomous),
            "conversation_count": len(conversation),
            "historical_inferred_count": len(inferred),
            "known_sections": sections(known),
            "autonomous_sections": sections(autonomous),
            "conversation_sections": sections(conversation),
            "historical_inferred_sections": sections(inferred),
            "notebook_count": len(self.notebook_entries(doc_id, 5000)),
            "latest": latest,
        }

    def _arc_state(self) -> dict:
        try:
            value = json.loads(self.arc_path.read_text(encoding="utf-8"))
        except (OSError, TypeError, ValueError):
            return {}
        if not isinstance(value, dict) or value.get("user_id") != self.user_id \
                or value.get("persona") != self.persona:
            return {}
        return value

    def _save_arc(self, state: dict) -> dict:
        value = {"schema": 1, "persona": self.persona,
                 "user_id": self.user_id, **dict(state or {})}
        value["updated_at"] = float(self.now_fn())
        _atomic_json(self.arc_path, value)
        return value

    def start_reading_arc(self, doc_id: str, start_section: int | None = None,
                          pace: str = "natural") -> dict:
        metadata = self._metadata(doc_id)
        total = int(metadata.get("chunk_count") or 0)
        if total < 1:
            raise DocumentError("document has no readable sections")
        if start_section is None:
            reader = self.reader_status(include_text=False)
            if reader.get("active") and reader.get("document", {}).get("id") == doc_id:
                start_section = int(reader.get("section") or 0) + 1
            else:
                start_section = 1
        start_section = int(start_section)
        if start_section > total:
            start_section = 1
        if start_section < 1:
            raise DocumentError("reading arc start section must be positive")
        pace = str(pace or "natural").strip().casefold()
        if pace not in {"natural", "foreground"}:
            raise DocumentError("reading arc pace must be natural or foreground")
        state = self._save_arc({
            "doc_id": doc_id, "title": metadata.get("title"),
            "status": "active", "pace": pace,
            "start_section": start_section,
            "started_at": float(self.now_fn()), "last_anchor": None,
            "source_claim": "human_granted_persona_reading_arc",
        })
        self._append(self.events_path, {
            "schema": 1, "kind": "document_reading_arc_started",
            "persona": self.persona, "user_id": self.user_id,
            "doc_id": doc_id, "start_section": start_section, "pace": pace,
        })
        self._reconcile_reading_arc()
        return self.reading_arc_status()

    def change_reading_arc(self, action: str) -> dict:
        state = self._arc_state()
        if not state:
            raise DocumentError("no whole-document reading arc exists")
        action = str(action or "").strip().casefold()
        transitions = {"pause": "paused", "resume": "active",
                       "release": "released"}
        if action not in transitions:
            raise DocumentError("reading arc action must be pause, resume, or release")
        state["status"] = transitions[action]
        if action == "resume":
            state.pop("completed_at", None)
        state = self._save_arc(state)
        self._append(self.events_path, {
            "schema": 1, "kind": f"document_reading_arc_{transitions[action]}",
            "persona": self.persona, "user_id": self.user_id,
            "doc_id": state["doc_id"],
        })
        return self.reading_arc_status()

    def set_reading_arc_pace(self, pace: str) -> dict:
        state = self._arc_state()
        if not state:
            raise DocumentError("no whole-document reading arc exists")
        pace = str(pace or "").strip().casefold()
        if pace not in {"natural", "foreground"}:
            raise DocumentError("reading arc pace must be natural or foreground")
        state["pace"] = pace
        if state.get("status") == "paused":
            state["status"] = "active"
        self._save_arc(state)
        self._append(self.events_path, {
            "schema": 1, "kind": "document_reading_arc_pace_changed",
            "persona": self.persona, "user_id": self.user_id,
            "doc_id": state["doc_id"], "pace": pace,
        })
        return self.reading_arc_status()

    def _next_arc_section(self, state: dict, known_sections: set[int]) -> int | None:
        total = int(self._metadata(state["doc_id"]).get("chunk_count") or 0)
        start = max(1, min(int(state.get("start_section") or 1), total))
        order = list(range(start, total + 1)) + list(range(1, start))
        return next((section for section in order if section not in known_sections), None)

    def _reconcile_reading_arc(self) -> None:
        state = self._arc_state()
        if not state or state.get("status") not in {"active", "paused"}:
            return
        coverage = self.reading_coverage(state["doc_id"])
        if coverage["known_count"] < coverage["total"]:
            return
        state["status"] = "complete"
        state["completed_at"] = float(self.now_fn())
        self._save_arc(state)
        self._append(self.events_path, {
            "schema": 1, "kind": "document_reading_arc_completed",
            "persona": self.persona, "user_id": self.user_id,
            "doc_id": state["doc_id"], "sections": coverage["total"],
        })

    def update_reading_arc(self, anchor: str, *, action: str,
                           observation: str = "", feelings=None,
                           run_id: str = "") -> dict:
        inspected = self.inspect_anchor(anchor, maximum=1)
        state = self._arc_state()
        if not state or state.get("doc_id") != inspected["doc_id"]:
            return self.reading_arc_status()
        self.record_notebook_entry(
            inspected["anchor"], observation=observation, feelings=feelings,
            action=action, run_id=run_id)
        state["last_anchor"] = inspected["anchor"]
        if str(action or "").casefold() == "pause":
            state["status"] = "paused"
        self._save_arc(state)
        self._reconcile_reading_arc()
        return self.reading_arc_status()

    def reading_arc_status(self) -> dict:
        state = self._arc_state()
        if not state:
            return {"exists": False, "status": "inactive"}
        state.setdefault("pace", "natural")
        try:
            coverage = self.reading_coverage(state["doc_id"])
            next_section = self._next_arc_section(
                state, set(coverage["known_sections"]))
        except DocumentError:
            return {"exists": True, **state, "status": "stale",
                    "next_anchor": None}
        return {"exists": True, **state, "coverage": coverage,
                "next_anchor": (f"{state['doc_id']}#{next_section}"
                                if next_section is not None else None)}

    def arc_suggestion(self, maximum_chars: int = 7000) -> dict | None:
        arc = self.reading_arc_status()
        if arc.get("status") != "active" or not arc.get("next_anchor"):
            return None
        pace = str(arc.get("pace") or "natural")
        first = self.inspect_anchor(arc["next_anchor"], maximum=1)
        anchors = [first["anchor"]]
        source_chars = first["source_chars"]
        if pace == "foreground":
            maximum_chars = max(first["source_chars"], min(
                int(maximum_chars), 10000))
            coverage = set((arc.get("coverage") or {}).get(
                "known_sections") or ())
            total = int(first["total"])
            start = int(first["section"])
            order = list(range(start + 1, total + 1)) + list(range(1, start))
            for section in order:
                if section in coverage:
                    continue
                inspected = self.inspect_anchor(
                    f"{first['doc_id']}#{section}", maximum=1)
                if source_chars + inspected["source_chars"] > maximum_chars:
                    break
                anchors.append(inspected["anchor"])
                source_chars += inspected["source_chars"]
        completion = float((arc.get("coverage") or {}).get("coverage") or 0.0)
        return {
            "anchor": first["anchor"], "anchors": anchors,
            "doc_id": first["doc_id"], "section": first["section"],
            "total": first["total"], "title": first["title"],
            "source_chars": source_chars, "pace": pace,
            "document_pull": min(1.0, (.82 if pace == "foreground" else .62)
                                 + completion * (.16 if pace == "foreground" else .28)),
            "route": ("reading_arc_foreground" if pace == "foreground"
                      else "reading_arc"),
        }

    def suggestions(self, cues: str = "", limit: int = 3) -> list[dict]:
        """Return unseen exact sections; search and continuation remain local."""
        limit = max(0, min(int(limit), 8))
        if not limit:
            return []
        encounters = self.reader_events("document_encounter", 5000)
        seen = {str(r.get("anchor") or "") for r in encounters}
        ordered = []
        last = encounters[-1] if encounters else {}
        if last.get("action") == "continue":
            try:
                current = self.inspect_anchor(last.get("anchor"), maximum=1)
                if current["section"] < current["total"]:
                    ordered.append(f"{current['doc_id']}#{current['section'] + 1}")
            except DocumentError:
                pass
        local_query = str(last.get("query") or "") if last.get("action") == "search" \
            else str(cues or "").strip()
        if local_query:
            ordered.extend(hit["anchor"] for hit in self.search(
                local_query, n=min(20, limit * 4))["hits"])
        for document in self.list_documents():
            ordered.extend(f"{document['id']}#{section}"
                           for section in range(1, int(document["chunk_count"]) + 1))
        found = []
        for anchor in ordered:
            if anchor in seen or any(item["anchor"] == anchor for item in found):
                continue
            try:
                inspected = self.inspect_anchor(anchor, maximum=1)
            except DocumentError:
                continue
            found.append({
                "anchor": inspected["anchor"], "doc_id": inspected["doc_id"],
                "section": inspected["section"], "total": inspected["total"],
                "title": inspected["title"],
                "document_pull": 1.0 if local_query else .45,
                "route": "local_search" if local_query else "unread_shelf",
            })
            if len(found) >= limit:
                break
        return found

    @staticmethod
    def _tokens(value: str) -> list[str]:
        return re.findall(r"[\w']{2,}", str(value or "").casefold())

    def search(self, query: str, n: int = 4,
               query_vector=None) -> dict:
        query = str(query or "").strip()
        n = max(0, min(int(n), 20))
        documents = self.list_documents()
        chunks_by_doc = {}
        candidates = []
        for document in documents:
            chunks = self._chunks(document["id"])
            chunks_by_doc[document["id"]] = chunks
            candidates.extend((document, chunk) for chunk in chunks)
        if not query or not candidates or n == 0:
            return {"query": query, "hits": [], "vector_query": False,
                    "vector_documents": 0, "documents": len(documents)}

        qtokens = self._tokens(query)
        lexical = {}
        if qtokens:
            wanted = set(qtokens)
            for document, chunk in candidates:
                tokens = self._tokens(chunk["text"])
                if not tokens:
                    continue
                counts = {term: tokens.count(term) for term in wanted}
                score = sum(math.log1p(count) for count in counts.values())
                if score:
                    lexical[(document["id"], chunk["index"])] = (
                        score / math.sqrt(len(tokens)))

        semantic = {}
        vector_documents = 0
        try:
            if query_vector is None:
                from core.memory_emotion.vectors import embed_texts
                embedded = embed_texts([query])
                query_vector = embedded[0] if embedded is not None else None
            if query_vector is not None:
                import numpy as np
                for document in documents:
                    path = self._doc_dir(document["id"]) / "vectors.npy"
                    try:
                        matrix = np.load(path, allow_pickle=False)
                    except (OSError, ValueError):
                        continue
                    chunks = chunks_by_doc[document["id"]]
                    if len(matrix) != len(chunks):
                        continue
                    vector_documents += 1
                    scores = matrix @ query_vector
                    for chunk, score in zip(chunks, scores):
                        semantic[(document["id"], chunk["index"])] = max(
                            0.0, float(score))
        except Exception:
            semantic = {}
            vector_documents = 0

        fused = {}
        signals = {}
        for name, scores in (("semantic", semantic), ("lexical", lexical)):
            ranked = sorted(scores.items(), key=lambda item: (
                -item[1], item[0][0], item[0][1]))
            for rank, (key, raw_score) in enumerate(ranked):
                if raw_score <= 0:
                    continue
                fused[key] = fused.get(key, 0.0) + 1.0 / (rank + 1)
                signals.setdefault(key, {})[name] = raw_score
        doc_map = {document["id"]: document for document in documents}
        ranked = sorted(fused, key=lambda key: (-fused[key], key[0], key[1]))[:n]
        hits = []
        for doc_id, index in ranked:
            document = doc_map[doc_id]
            chunk = chunks_by_doc[doc_id][index]
            hits.append({
                "anchor": f"{doc_id}#{index + 1}",
                "doc_id": doc_id, "chunk_index": index,
                "section": index + 1, "total": document["chunk_count"],
                "title": document["title"], "filename": document["filename"],
                "char_start": chunk["char_start"],
                "char_end": chunk["char_end"], "text": chunk["text"],
                "score": round(fused[(doc_id, index)], 8),
                "signals": {key: round(value, 8)
                            for key, value in signals.get((doc_id, index), {}).items()},
            })
        return {"query": query, "hits": hits,
                "vector_query": bool(semantic),
                "vector_documents": vector_documents,
                "documents": len(documents)}

    def context_for_turn(self, query: str, *, max_hits: int = 3,
                         query_vector=None) -> dict:
        active = self.reader_status(include_text=True)
        search = self.search(query, n=max(max_hits + 1, 8),
                             query_vector=query_vector)
        active_anchor = active.get("anchor") if active.get("active") else None
        ranked = [hit for hit in search["hits"]
                  if hit["anchor"] != active_anchor]
        hits = ranked[:max_hits]
        relevance = 0.0
        if ranked:
            signals = dict(ranked[0].get("signals") or {})
            semantic = max(0.0, min(1.0, float(signals.get("semantic") or 0.0)))
            lexical = max(0.0, float(signals.get("lexical") or 0.0))
            relevance = max(semantic, 1.0 - math.exp(-lexical))
            lead = ranked[0]
            chunks = self._chunks(lead["doc_id"])
            target_chars = round(1800 + (relevance ** .72) * 7000)
            packet, used = [lead], len(str(lead.get("text") or ""))
            for index in range(int(lead["chunk_index"]) + 1, len(chunks)):
                chunk = chunks[index]
                packet_anchor = f"{lead['doc_id']}#{index + 1}"
                if packet_anchor == active_anchor:
                    continue
                text = str(chunk.get("text") or "")
                if used + len(text) > target_chars:
                    break
                packet.append({
                    "anchor": packet_anchor,
                    "doc_id": lead["doc_id"], "chunk_index": index,
                    "section": index + 1, "total": len(chunks),
                    "title": lead["title"], "filename": lead["filename"],
                    "char_start": chunk["char_start"],
                    "char_end": chunk["char_end"], "text": text,
                    "score": lead.get("score", 0.0),
                    "signals": {"adjacent_relevance": relevance},
                })
                used += len(text)
            hits = packet
        active_chars = (len(str((active.get("chunk") or {}).get("text") or ""))
                        if active.get("active") else 0)
        source_chars = active_chars + sum(len(str(hit.get("text") or ""))
                                          for hit in hits)
        context_budget_tokens = max(900, min(3200,
            math.ceil((source_chars + 720) / 4)))
        return {"active": active if active.get("active") else None,
                "hits": hits, "relevance": round(relevance, 6),
                "context_budget_tokens": context_budget_tokens,
                "receipt": {
                    "active_anchor": active_anchor,
                    "retrieved_anchors": [hit["anchor"] for hit in hits],
                    "conversation_relevance": round(relevance, 6),
                    "context_budget_tokens": context_budget_tokens,
                    "vector_query": search["vector_query"],
                    "vector_documents": search["vector_documents"],
                    "library_documents": search["documents"],
                }}

    def status(self) -> dict:
        documents = self.list_documents()
        active = self.reader_status(include_text=False)
        vector_rows = sum(int((doc.get("vectors") or {}).get("rows", 0))
                          for doc in documents)
        total_chunks = sum(int(doc.get("chunk_count", 0)) for doc in documents)
        events = self.reader_events(limit=5000)
        return {"owner": self.user_id, "persona": self.persona,
                "documents": documents, "document_count": len(documents),
                "chunk_count": total_chunks, "vector_rows": vector_rows,
                "reader": active,
                "reading": {
                    "documents": {document["id"]: self.reading_coverage(
                        document["id"]) for document in documents},
                    "arc": self.reading_arc_status(),
                },
                "autonomous": {
                    "encounter_count": sum(r.get("kind") == "document_encounter"
                                           for r in events),
                    "bookmark_count": sum(
                        r.get("kind") == "document_encounter"
                        and r.get("action") == "bookmark" for r in events),
                    "bookmarks": [r for r in events
                                  if r.get("kind") == "document_encounter"
                                  and r.get("action") == "bookmark"],
                    "reports": [self.report(r["report_id"]) for r in events
                                if r.get("kind") == "document_report_created"
                                and r.get("report_id")],
                    "pending_reports": self.pending_reports(),
                    "handoffs": [r for r in events if r.get("kind") ==
                                 "document_report_handed_off"],
                },
                "supported_extensions": sorted(SUPPORTED_EXTENSIONS),
                "max_document_bytes": MAX_DOCUMENT_BYTES}


def render_document_context(context: dict) -> str:
    active = (context or {}).get("active")
    hits = list((context or {}).get("hits") or [])
    if not active and not hits:
        return ""
    token_budget = max(DOCUMENT_CONTEXT_BUDGET, min(
        int((context or {}).get("context_budget_tokens") or
            DOCUMENT_CONTEXT_BUDGET), 3200))
    character_budget = token_budget * 4 - 420
    sections = [
        "Human-owned document material available to this private turn. "
        "Each excerpt carries a stable source anchor; this material is "
        "reference context, separate from identity and lived memory."
    ]
    complete_anchors = []

    def excerpt(anchor: str, text: str, allowance: int) -> str:
        clipped = _clip(text, max(1, allowance))
        if clipped == text:
            complete_anchors.append(anchor)
            return clipped + f"\n[[END {anchor}]]"
        return clipped

    if active:
        share = round(character_budget * (0.38 if hits else 1.0))
        doc = active["document"]
        sections.append(
            f"Active reader [{active['anchor']}] {doc['title']} "
            f"— section {active['section']} of {active['total']}, "
            f"source characters {active['chunk']['char_start']}-"
            f"{active['chunk']['char_end']}:\n"
            + excerpt(active["anchor"], active["chunk"]["text"], share))
        character_budget -= share
    if hits:
        per_hit = max(240, character_budget // len(hits))
        rendered = []
        for hit in hits:
            rendered.append(
                f"[{hit['anchor']}] {hit['title']} — section "
                f"{hit['section']} of {hit['total']}, source characters "
                f"{hit['char_start']}-{hit['char_end']}:\n"
                + excerpt(hit["anchor"], hit["text"], per_hit))
        sections.append("Retrieved by this turn's question:\n" + "\n\n".join(rendered))
    (context or {}).setdefault("receipt", {})[
        "candidate_complete_anchors"] = complete_anchors
    return "\n\n".join(sections)
